"""
Document preprocessing utilities for the building information extractor
Handles conversion from various formats to plain text
"""

import subprocess
from pathlib import Path
from typing import Optional


class DocumentConverter:
    """Convert various document formats to plain text"""
    
    @staticmethod
    def docx_to_text(docx_path: str) -> str:
        """Convert DOCX to plain text using python-docx"""
        try:
            from docx import Document
            
            doc = Document(docx_path)
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            return '\n\n'.join(text_parts)
        except ImportError:
            raise ImportError("python-docx not installed. Install with: pip install python-docx")
    
    @staticmethod
    def pdf_to_text(pdf_path: str) -> str:
        """Convert PDF to plain text using PyMuPDF (fitz)"""
        try:
            import fitz
            
            text_parts = []
            doc = fitz.open(pdf_path)
            for page in doc:
                # Use block-level extraction and sort by reading order (y, then x)
                blocks = page.get_text("blocks")
                blocks.sort(key=lambda b: (b[1], b[0]))
                
                # Filter out small irrelevant blocks and extract text
                text = "\n".join([b[4] for b in blocks if len(b) >= 5])
                if text.strip():
                    text_parts.append(text)
            
            return '\n\n'.join(text_parts)
        except ImportError:
            raise ImportError("PyMuPDF not installed. Install with: pip install pymupdf")
    
    @staticmethod
    def txt_to_text(txt_path: str) -> str:
        """Read plain text file"""
        with open(txt_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    @staticmethod
    def convert_to_text(file_path: str) -> str:
        """
        Auto-detect file type and convert to text
        
        Args:
            file_path: Path to the document file
        
        Returns:
            Plain text content
        """
        path = Path(file_path)
        extension = path.suffix.lower()
        
        converters = {
            '.txt': DocumentConverter.txt_to_text,
            '.docx': DocumentConverter.docx_to_text,
            '.pdf': DocumentConverter.pdf_to_text,
        }
        
        if extension not in converters:
            raise ValueError(f"Unsupported file type: {extension}. Supported: {list(converters.keys())}")
        
        return converters[extension](file_path)
    
    @staticmethod
    def preprocess_text(text: str) -> str:
        """
        Preprocess text to improve extraction quality
        
        Args:
            text: Raw text
        
        Returns:
            Preprocessed text
        """
        # Remove excessive whitespace
        text = '\n'.join(line.strip() for line in text.split('\n'))
        
        # Remove multiple consecutive blank lines
        while '\n\n\n' in text:
            text = text.replace('\n\n\n', '\n\n')
        
        # Remove page numbers (common pattern: just a number on a line)
        lines = text.split('\n')
        cleaned_lines = []
        for line in lines:
            # Skip lines that are just numbers (likely page numbers)
            if line.strip().isdigit() and len(line.strip()) <= 3:
                continue
            cleaned_lines.append(line)
        
        text = '\n'.join(cleaned_lines)
        
        return text.strip()


def convert_document(input_path: str, output_path: Optional[str] = None) -> str:
    """
    Convert a document to text and optionally save to file
    
    Args:
        input_path: Path to input document
        output_path: Optional path to save text output
    
    Returns:
        Extracted text
    """
    converter = DocumentConverter()
    
    print(f"Converting {input_path}...")
    text = converter.convert_to_text(input_path)
    
    print(f"Preprocessing text...")
    text = converter.preprocess_text(text)
    
    if output_path:
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text)
        print(f"✓ Text saved to {output_path}")
    
    print(f"✓ Extracted {len(text)} characters")
    return text


if __name__ == "__main__":
    import sys
    
    if len(sys.argv) < 2:
        print("Usage: python document_converter.py <input_file> [output_text_file]")
        print("\nSupported formats: .txt, .docx, .pdf")
        print("\nExample:")
        print("  python document_converter.py paper.pdf paper.txt")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None
    
    if not output_file:
        # Auto-generate output filename
        input_path = Path(input_file)
        output_file = input_path.stem + '_converted.txt'
    
    text = convert_document(input_file, output_file)
